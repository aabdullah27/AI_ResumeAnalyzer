import streamlit as st
import os
import tempfile
import faiss
from llama_index.core import VectorStoreIndex, Document
from llama_index.vector_stores.faiss import FaissVectorStore
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.groq import Groq
from llama_index.core import Settings
from pypdf import PdfReader
from docx import Document as DocxDocument


def extract_text(uploaded_file):
    """Extract text from the uploaded resume file."""
    # Save uploaded file temporarily
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as temp_file:
        temp_file.write(uploaded_file.getvalue())
        temp_file_path = temp_file.name

    text = ""
    try:
        if uploaded_file.type == "application/pdf":
            pdf_reader = PdfReader(temp_file_path)
            text = "".join(page.extract_text() for page in pdf_reader.pages)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = DocxDocument(temp_file_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        else:  # Assume it's a text file
            with open(temp_file_path, 'r') as file:
                text = file.read()
    finally:
        os.unlink(temp_file_path)  # Clean up the temporary file

    return text


def create_query():
    """Construct the query for the LLM."""
    return '''
    Based on this resume or CV, perform the following:
    1. Experience Analysis: Determine the total years of experience and categorize the candidate as Junior, Mid-level, or Senior.
    2. Extract top 3 projects or accomplishments listed in the resume.
    3. ATS Compatibility Score: Provide a score out of 10 based on formatting, keyword usage, and structure.
    4. Readability and Clarity Rating: Rate the clarity of the resume on a scale of 1 to 10.
    5. Job Roles: Tell the top 3 JOB roles this candidate is perfect for in a numbering manner.
    6. Name and Age: Show the name and age if available.
    NOTE: If there is any content except Resume or CV, provide a polite answer to upload a CV for analysis.
    '''


# Streamlit app setup
st.set_page_config(page_title="Resume Analyzer", page_icon="📝",layout="wide")
st.title("Resume Analyzer 📝")

# Sidebar configuration
with st.sidebar:
    st.header("Configuration ⚙️")
    groq_api_key = st.text_input("Enter your Groq API Key 🔑", type="password")
    uploaded_file = st.file_uploader("Upload your resume (PDF, DOCX, or TXT) 📄", type=["pdf", "docx", "txt"])
    analyze_button = st.button("Analyze Resume 🔍")

# Main content area
if not groq_api_key or not uploaded_file:
    st.warning("⚠️ Please enter your Groq API key and upload a resume.")
elif analyze_button:
    with st.spinner("Analyzing your resume... ⏳"):
        try:
            # Extract text from the uploaded file
            text = extract_text(uploaded_file)

            # Set up RAG components
            embed_model = HuggingFaceEmbedding(model_name="BAAI/bge-small-en")
            Settings.embed_model = embed_model  # Set the embedding model in the Settings

            # Create a FAISS index
            d = 768  # Dimension (you can adjust this based on your embedding model)
            faiss_index = faiss.IndexFlatL2(d)

            # Create the FAISS vector store
            vector_store = FaissVectorStore(faiss_index=faiss_index)

            # Initialize the Groq model with an appropriate model name
            groq_model_name = "mixtral-8x7b-32768"  # Replace with the actual model name you want to use
            llm = Groq(api_key=groq_api_key, model=groq_model_name)

            # Set the LLM in Settings
            Settings.llm = llm

            # Create documents and index
            documents = [Document(text=text)]
            index = VectorStoreIndex.from_documents(documents, vector_store=vector_store)

            # Initialize the query engine
            query_engine = index.as_query_engine()

            # Execute the query
            query = create_query()
            response = query_engine.query(query)

            # Display the response
            if not response.response:
                st.warning(
                    "⚠️ No relevant information found in the resume. Please ensure the document is a valid CV or resume.")
            else:
                st.markdown("### Resume Analysis Results 📊")
                st.write(response.response)

        except Exception as e:
            st.error(f"An error occurred during analysis: {e}")

st.sidebar.markdown("---")
st.sidebar.markdown("Created with Streamlit, LlamaIndex, FAISS, and Groq 💡 By ABDULLAH 🗿")
